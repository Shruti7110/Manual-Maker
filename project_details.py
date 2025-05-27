from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def insert_project_details(doc, project_info):
    if doc is None or not hasattr(doc, 'paragraphs'):
        return None, False
        
    if not isinstance(project_info, dict):
        return doc, False
    """
    Insert project details at the specified placeholder.
    
    Args:
        doc (Document): Document object
        project_info (dict): Dictionary containing project information
    
    Returns:
        Document: Updated document
        bool: Whether placeholder was found and replaced
    """
    placeholder_found = False
    
    # Find the placeholder
    for i, para in enumerate(doc.paragraphs):
        if "{{project_details}}" in para.text:
            # Mark as found
            placeholder_found = True
            
            # Remove the placeholder text
            para.text = para.text.replace("{{Project Details}}", "")
            
            # Add formatted project details
            para.add_run("Project Details:").bold = True
            para.add_run("\n\n")
            
            # Add each project detail
            for key, value in project_info.items():
                if value and value != "NOT FOUND":
                    formatted_key = key.replace("_", " ").title()
                    para.add_run(f"{formatted_key}: ").bold = True
                    para.add_run(f"{value}\n")
            
            break
    
    return doc, placeholder_found

def insert_electrical_specifications(doc, electrical_specs):
    """
    Insert electrical specifications at the specified placeholder.
    
    Args:
        doc (Document): Document object
        electrical_specs (dict): Dictionary containing electrical specifications
    
    Returns:
        Document: Updated document
        bool: Whether placeholder was found and replaced
    """
    placeholder_found = False
    
    # Find the placeholder
    for i, para in enumerate(doc.paragraphs):
        if "{{electrical_specs}}" in para.text:
            # Mark as found
            placeholder_found = True
            
            # Remove the placeholder text
            para.text = para.text.replace("{{Electrical_Specifications}}", "")
            
            # Add heading
            para.add_run("Electrical Specifications:").bold = True
            para.add_run("\n\n")
            
            # Add each specification
            has_specs = False
            for key, value in electrical_specs.items():
                if value:  # Only add non-empty values
                    has_specs = True
                    formatted_key = key.replace("_", " ").title()
                    para.add_run(f"{formatted_key}: ").bold = True
                    para.add_run(f"{value}\n")
            
            # If no specifications were provided
            if not has_specs:
                para.add_run("No electrical specifications provided.")
            
            break
    
    return doc, placeholder_found

def extract_project_info(txt_path, info):
    """
    Extracts project_name, customer, and project_no from a TXT file.
    """
    
    with open(txt_path, "r") as file:
        for line in file:
            for key in info.keys():
                if key in line:
                    # Split by '=' and strip trailing commas/spaces
                    value = line.split("=")[1].strip().rstrip(",")
                    info[key] = value
    return info


def insert_project_info(docx_path):
    """
    Inserts the extracted info into a DOCX file at the placeholder location.
    """
    doc = Document(docx_path)
    output_path = docx_path
    placeholder = "{{Project Details}}"
    # Initialize info dictionary
    info = {"project_name": "", "customer": "", "project_no": ""}
    project_info = extract_project_info(txt_path="uploads/Project_info/Project_info.txt", info=info)

    # Create content block
    content = (
    f"Project Name = {project_info['project_name']}, \n"
    f"Customer = {project_info['customer']}, \n"
    f"Project No = {project_info['project_no']}"
)
    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = para.text.replace(placeholder, "")
            run = para.add_run(content)
            break

    doc.save(output_path)
    print(f"✅ Updated project info in DOCX saved at: {output_path}")
    
    return doc, True

def insert_machine_specifications(docx_path):
    """
    Inserts the extracted info into a DOCX file at the placeholder location.
    """
    doc = Document(docx_path)
    info = {"machine_specs": ""}
    output_path = docx_path
    placeholder = "{{Machine_Specifications}}"
    project_info = extract_project_info(txt_path="uploads/Project_info/Project_info.txt", info=info)
    
    # Create content block
    content = f"{project_info['machine_specs']}"
    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = para.text.replace(placeholder, "")
            run = para.add_run(content)
            break

    doc.save(output_path)
    print(f"✅ Updated machine specs in DOCX saved at: {output_path}")
    return doc, True

def insert_electrical_specifications(docx_path):
    """
    Inserts the extracted info into a DOCX file at the placeholder location.
    """
    doc = Document(docx_path)
    info = {"Voltage": "", "Power": "", "Current": "", "Frequency": ""}
    output_path = docx_path
    
    placeholder = "{{Electrical_Specifications}}"
    project_info = extract_project_info(txt_path="uploads/Project_info/Project_info.txt", info=info)
    
    # Create content block
    content = (
        f"Voltage:{project_info['Voltage']},   "
        f"Power:{project_info['Power']},   "
        f"Current:{project_info['Current']},   "
        f"Frequency:{project_info['Frequency']}"
    )
    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = para.text.replace(placeholder, "")
            run = para.add_run(content)
            break

    doc.save(output_path)
    print(f"✅ Updated electrical specs in DOCX saved at: {output_path}")
    
    return doc, True


