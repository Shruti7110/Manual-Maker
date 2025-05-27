import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import tempfile
import PyPDF2
from pdf2image import convert_from_path
from docx import Document
import fitz  # PyMuPDF
from PIL import Image
import io
import glob


def pdf_to_images(pdf_path, output_dir):
    """
    Converts each page of the PDF into a separate image and saves them in the output directory.

    Args:
        pdf_path (str): Path to the PDF file.
        output_dir (str): Directory to save the output images.
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF not found: {pdf_path}")
    
    os.makedirs(output_dir, exist_ok=True)

    # Convert PDF to images
    print(f"🔄 Converting PDF pages to images from: {pdf_path}")
    images = convert_from_path(pdf_path, dpi=300)

    for i, img in enumerate(images):
        image_path = os.path.join(output_dir, f"PDF_page_{i+1}.png")
        img.save(image_path, "PNG")
        print(f"✅ Saved: {image_path}")

    print("📁 All pages saved as images.")

def insert_pdf_images_at_placeholder(docx_path, output_dir, placeholder):
    """
    Inserts all images from image_dir into the docx file at the placeholder location.

    Args:
        docx_path (str): Path to the .docx file (will be updated in-place).
        image_dir (str): Directory containing image files to insert.
        placeholder (str): Placeholder text to be replaced with images.
    """
    doc = Document(docx_path)
    image_width = 6.0
    # Sort image files (natural order)
    image_files = sorted(
        [f for f in os.listdir(output_dir) if f.lower().endswith((".png", ".jpg", ".jpeg"))],
        key=lambda x: int("".join(filter(str.isdigit, x)) or 0)
    )

    if not image_files:
        print("⚠️ No image files found to insert.")
        return

    # Locate placeholder
    for i, para in enumerate(doc.paragraphs):
        if placeholder in para.text:
            placeholder_para = para
            break
    else:
        raise ValueError(f"❌ Placeholder '{placeholder}' not found in the document.")

    # Clear placeholder text
    placeholder_para.text = placeholder_para.text.replace(placeholder, "").strip()

    def insert_paragraph_after(paragraph: Paragraph):
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        return Paragraph(new_p, paragraph._parent)

    # Insert images after placeholder
    current_para = placeholder_para
    for img_name in image_files:
        img_path = os.path.join(output_dir, img_name)

        new_para = insert_paragraph_after(current_para)
        run = new_para.add_run()
        run.add_picture(img_path, width=Inches(image_width))
        new_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Optional: add break
        current_para = insert_paragraph_after(new_para)
        current_para.add_run().add_break()

    doc.save(docx_path)
    print(f"✅ Updated DOCX saved: {docx_path}")

def process_eplan_pdf_to_docx(folder_path, template_path):
    """
    Extracts headings and images from a EPLAN PowerPoint and inserts them into a DOCX template.

    Args:
        folder_path (str): Path to the folder containing the EPLAN PowerPoint file.
        template_path (str): Path to the base DOCX template.
        output_docx_path (str): Path to save the final generated manual.
        placeholder (str): Placeholder text in the DOCX to be replaced.
    """
    pdf_files = glob.glob(os.path.join(folder_path, "*.pdf"))
    if len(pdf_files) == 0:
        print("⚠️ No PDF file found in EPLAN folder. Aborting.")
        return
    elif len(pdf_files) > 1:
        print("⚠️ Multiple PDF files found in EPLAN folder. Aborting.")
        return

    pdf_path = pdf_files[0]
    # print(f"📂 Found PDF file: {pdf_path}")

    output_dir = os.path.abspath("eplan_img_extracted")
    placeholder= "{{Upload_Electrical_drawing_here}}"

    # Step 1: Extract slide data (images + headings via OCR)
    print("Extracting EPLAN slide data...")
    # pdf_to_images(pdf_path, output_dir)

    print("Inserting EPLAN slide data into DOCX...")
    # Step 2: Insert extracted data into DOCX
    insert_pdf_images_at_placeholder(
        docx_path = template_path, 
        output_dir = output_dir, 
        placeholder = placeholder)

def process_alarms_pdf_to_docx(folder_path, template_path):
    """
    Extracts headings and images from a Alarms pdf and inserts them into a DOCX template.

    Args:
        folder_path (str): Path to the folder containing the Alarms pdf file.
        template_path (str): Path to the base DOCX template.
        output_docx_path (str): Path to save the final generated manual.
        placeholder (str): Placeholder text in the DOCX to be replaced.
    """
    pdf_files = glob.glob(os.path.join(folder_path, "*.pdf"))
    if len(pdf_files) == 0:
        print("⚠️ No PDF file found in Alarms folder. Aborting.")
        return
    elif len(pdf_files) > 1:
        print("⚠️ Multiple PDF files found in Alarms folder. Aborting.")
        return

    pdf_path = pdf_files[0]
    # print(f"📂 Found PDF file: {pdf_path}")

    output_dir = os.path.abspath("alarms_img_extracted")
    placeholder= "{{Upload_alarms_doc_here}}"

    # Step 1: Extract slide data (images + headings via OCR)
    print("Extracting Alarms slide data...")
    pdf_to_images(pdf_path, output_dir)

    print("Inserting Alarms slide data into DOCX...")
    # Step 2: Insert extracted data into DOCX
    insert_pdf_images_at_placeholder(
        docx_path = template_path, 
        output_dir = output_dir, 
        placeholder = placeholder)
