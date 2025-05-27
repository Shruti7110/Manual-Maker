import streamlit as st
import os
import base64
from datetime import datetime
import tempfile
from docx import Document

def save_template():
    """
    Save the uploaded template or create a basic one if none is provided
    """
    # Create template directory if it doesn't exist
    if not os.path.exists("template"):
        os.makedirs("template")
    
    # Path to save the template
    template_path = os.path.join("template", "manual_template.docx")
    
    # Create a basic template with proper structure
    doc = Document()
    
    # Add title
    doc.add_heading("OPERATIONS & INSTALLATION MANUAL", 0)
    
    # Add placeholder for machine photo
    doc.add_paragraph("Insert Machine Photo Here")
    
    # Add basic information section
    basic_info = doc.add_paragraph()
    basic_info.add_run("Customer   : ").bold = True
    basic_info.add_run("\nProject No : ").bold = True
    basic_info.add_run("\nDate       : ").bold = True
    
    # Add project name
    doc.add_heading("PROJECT NAME", 1)
    
    # Add ToC heading
    doc.add_heading("TABLE OF CONTENTS", 1)
    
    # Add main sections
    doc.add_heading("1. Installation Guide", 1)
    doc.add_heading("1.1 User", 2)
    doc.add_heading("1.2 Machine Safety", 2)
    doc.add_heading("1.3 Operator Safety", 2)
    
    doc.add_heading("2. About the Machine", 1)
    doc.add_heading("2.1 Machine Specifications", 2)
    doc.add_heading("2.2 Machine Overview", 2)
    doc.add_heading("2.3 Sequence Of Operation", 2)
    
    doc.add_heading("3. Machine Power On", 1)
    
    doc.add_heading("4. HMI", 1)
    doc.add_heading("4.1 HMI Login", 2)
    doc.add_heading("4.2 HMI Screens", 2)
    
    # Save the document
    doc.save(template_path)
    print(f"Template saved at: {template_path}")
    
    return template_path

if __name__ == "__main__":
    template_path = save_template()
    print(f"Template created at: {template_path}")