# Fixes the None check in the process_machine_photos function to prevent errors when doc is None or doesn't have paragraphs.
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os


def insert_images_with_placeholder(template_doc, placeholder, image_paths):
    """
    Insert images at a specific placeholder in a Word document.

    Args:
        template_doc (Document): Document object
        placeholder (str): Placeholder text
        image_paths (list): List of image paths

    Returns:
        Document: Updated document
        bool: Whether placeholder was found and replaced
    """
    placeholder_found = False

    # Find the placeholder
    for i, para in enumerate(template_doc.paragraphs):
        if placeholder in para.text:
            # Mark placeholder as found
            placeholder_found = True

            # Remove the placeholder text
            para.text = para.text.replace(placeholder, "")

            # Insert images
            for j, img_path in enumerate(image_paths):
                try:
                    if j == 0:
                        # For first image, add directly to current paragraph
                        run = para.add_run()
                        run.add_picture(img_path, width=Inches(6))
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    else:
                        # For additional images, add a new paragraph
                        new_para = template_doc.add_paragraph()
                        run = new_para.add_run()
                        run.add_picture(img_path, width=Inches(6))
                        new_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        # Add caption
                        caption_para = template_doc.add_paragraph(f"Image {j+1}")
                        caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                except Exception as e:
                    print(f"Error adding image {img_path}: {str(e)}")
            break

    return template_doc, placeholder_found

def get_images_from_folder(folder_path):
    """
    Get all image files from a folder.

    Args:
        folder_path (str): Path to the folder

    Returns:
        list: List of image file paths
    """
    image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp']
    image_paths = []

    if os.path.exists(folder_path):
        for filename in os.listdir(folder_path):
            # Check if the file is an image
            if any(filename.lower().endswith(ext) for ext in image_extensions):
                image_paths.append(os.path.join(folder_path, filename))

    return image_paths

def process_machine_photos(docx_path, output_dir=None):
    """
    Process machine photos and insert them into the document.

    Args:
        doc (Document): Document object
        output_dir (str): Output directory for artifacts

    Returns:
        Document: Updated document
        bool: Whether photos were processed
    """
    doc= Document(docx_path)
    if doc is None or not hasattr(doc, 'paragraphs'):
        return None, False

    machine_photos_folder = "uploads/Machine_Photos"
    image_paths = get_images_from_folder(machine_photos_folder)

    if not image_paths:
        return doc, False

    # Insert machine photos at Cover Photo placeholder
    doc, placeholder_found = insert_images_with_placeholder(
        doc, "{{Insert_Cover_Photo_Here}}", image_paths
    )
    print("Inserted machine photos into the document.")
    doc.save(docx_path)
    
    return doc, True

def process_layout_photos(docx_path, output_dir=None):
    """
    Process layout photos and insert them into the document.

    Args:
        doc (Document): Document object
        output_dir (str): Output directory for artifacts

    Returns:
        Document: Updated document
        bool: Whether photos were processed
    """
    doc= Document(docx_path)
    layout_photos_folder = "uploads/Layout_Photos"
    image_paths = get_images_from_folder(layout_photos_folder)

    if not image_paths:
        return doc, False

    # Insert layout photos at Layout placeholder
    doc, placeholder_found = insert_images_with_placeholder(
        doc, "{{Upload_Machine_Layout_here}}", image_paths
    )
    print("Inserted layout photos into the document.")
    doc.save(docx_path)
    return doc, True

def process_pneumatic_photos(docx_path, output_dir=None):
    """
    Process pneumatic photos and insert them into the document.

    Args:
        doc (Document): Document object
        output_dir (str): Output directory for artifacts

    Returns:
        Document: Updated document
        bool: Whether photos were processed
    """
    doc= Document(docx_path)
    pneumatic_photos_folder = "uploads/Pneumatic"
    image_paths = get_images_from_folder(pneumatic_photos_folder)

    if not image_paths:
        return doc, False

    # Insert pneumatic photos at Pneumatic placeholder
    doc, placeholder_found = insert_images_with_placeholder(
        doc, "{{Upload_Pneumatic_Circuit_Here}}", image_paths
    )
    print("Inserted pneumatic photos into the document.")
    doc.save(docx_path)
    return doc, True
   

def remove_unused_placeholders(docx_path):
    """
    Remove any unused placeholders from the document.

    Args:
        doc (Document): Document object

    Returns:
        Document: Updated document
    """
    doc= Document(docx_path)
    placeholders = [
        "{{machine_photo}}",
        "{{Upload_Machine_Layout_here}}",
        "{{MACHINE_OVERVIEW_DAP}}",
        "{{Upload_SOP_here}}",
        "{{Upload_HMI_here}}",
        "{{Upload_scada_screens_here}}",
        "{{Electrical_Specifications}}",#
        "{{Upload_Pneumatic_Circuit_Here}}",
        "{{project_details}}",
        "{{Machine_Specifications}}",#
        "{{HMI_SLIDES}}",#
        "{{Upload_alarms_doc_here}}",
        "{{Upload_MBOM}}",
        "{{Upload_Electrical_drawing_here}}",
        "{{Upload_other_docs_here}}"
    ]

    for i, para in enumerate(doc.paragraphs):
        for placeholder in placeholders:
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, "")

    print("Removed unused placeholders from the document.")
    doc.save(docx_path)
    return doc